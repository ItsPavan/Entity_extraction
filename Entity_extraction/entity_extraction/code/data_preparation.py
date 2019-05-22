import os
from docx import Document
import pandas as pd
import numpy as np
import re
from nltk.tokenize import RegexpTokenizer
# from nltk.tokenize import word_tokenize
from nltk import pos_tag
tokenizer = RegexpTokenizer(r'\w+')
data_path_train = 'A:\\zycus-test\\workspace-alewo4\\data\\training'
data_path_test = 'A:\\zycus-test\\workspace-alewo4\\data\\eval'
files = [file for file in os.listdir(data_path_test)]
documents_list = []
for f in range(0,len(files)):
    full_text = []
    document = Document(data_path_test+'/'+files[f])
    # print (document.paragraphs)
    for para in document.paragraphs:
        l = str(para.text.encode('utf-8').strip())[1:]
        l = l.strip('\'')
        l = 'SOS'+' '+l
        full_text.append(l)
        full_text = [line for line in full_text if line != "''"]
    documents_list.append(full_text)
r,s,pos = [],[],[]
counter = 0
for i in documents_list:
    for j in i:
        counter += 1
        raw = tokenizer.tokenize(j)
        # raw = word_tokenize(j)
        r.append(raw)
        s.append('sentence'+str(counter))

df_path = "A:\\zycus-test\\workspace-alewo4\\data"
'''
df = pd.DataFrame({'sentence':s,'word':r})
# split one row to multiple rows in a column by preserving another column in pandas dataframe
df = pd.DataFrame({col:np.repeat(df['sentence'].values, df['word'].str.len())for col in df.columns.difference(['word'])}).assign(**{'word':np.concatenate(df['word'].values)})[df.columns.tolist()]
df['tag'] = 'O'
# print (df.tail(5))

for i in pos_tag(df['word']):
    pos.append(i[1])
df['POS'] = pos
df.to_csv (r'A:\\zycus-test\\workspace-alewo4\\data\\train.csv', index = None, header=True)

tag = []
df = pd.read_csv(df_path+"\\"+"train.csv")
t = df['tag']
tag.append(t[1])
print(tag)
'''
#comment when writing to train file
'''
df_test = pd.DataFrame({'sentence':s,'word':r})
# split one row to multiple rows in a column by preserving another column in pandas dataframe
df_test = pd.DataFrame({col:np.repeat(df_test['sentence'].values, df_test['word'].str.len())for col in df_test.columns.difference(['word'])}).assign(**{'word':np.concatenate(df_test['word'].values)})[df_test.columns.tolist()]

# print (df.tail(5))
df_test.to_csv (r'A:\\zycus-test\\workspace-alewo4\\data\\test.csv', index = None, header=True)
df_path = "A:\\zycus-test\\workspace-alewo4\\data"
df_test = pd.read_csv(df_path+"\\"+"test.csv")
for i in pos_tag(df_test['word']):
    pos.append(i[1])
df_test['POS'] = pos
df_test['tag'] = 'O'
df_test.to_csv (r'A:\\zycus-test\\workspace-alewo4\\data\\test.csv', index = None, header=True)

'''